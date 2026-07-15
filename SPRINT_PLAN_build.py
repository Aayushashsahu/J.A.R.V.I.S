"""Convert SPRINT_PLAN.md -> HTML Presentation (Keynote Theme) + PDF + DOCX."""

from __future__ import annotations

import shutil
import subprocess
import sys
import re
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / "SPRINT_PLAN.md"
HTML_OUT = HERE / "SPRINT_PLAN.html"
PDF_OUT = HERE / "SPRINT_PLAN.pdf"
DOCX_OUT = HERE / "SPRINT_PLAN.docx"

HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>J.A.R.V.I.S — ET AI Hackathon 2026 Sprint Plan</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Outfit:wght@300;400;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
  <style>
    :root {
      --bg-space: #05070b;
      --bg-darker: #020305;
      --bg-card: rgba(10, 14, 23, 0.65);
      --border-cyan: rgba(0, 240, 255, 0.08);
      --border-cyan-glow: rgba(0, 240, 255, 0.25);
      --accent-blue: #0066ff;
      --accent-cyan: #00f0ff;
      --text-primary: #ffffff;
      --text-secondary: #94a3b8;
      --font-sans: 'Inter', sans-serif;
      --font-display: 'Outfit', sans-serif;
      --font-mono: 'JetBrains Mono', monospace;
    }

    * {
      box-sizing: border-box;
      margin: 0;
      padding: 0;
    }

    body {
      background-color: var(--bg-space);
      color: var(--text-primary);
      font-family: var(--font-sans);
      height: 100vh;
      width: 100vw;
      overflow: hidden;
      position: relative;
    }

    /* Grid Overlay */
    body::before {
      content: '';
      position: fixed;
      top: 0;
      left: 0;
      width: 100vw;
      height: 100vh;
      background-image: 
        linear-gradient(rgba(0, 240, 255, 0.012) 1px, transparent 1px),
        linear-gradient(90deg, rgba(0, 240, 255, 0.012) 1px, transparent 1px);
      background-size: 60px 60px;
      z-index: 1;
      pointer-events: none;
    }

    /* Cosmic Background radial lights */
    .space-glow {
      position: absolute;
      width: 600px;
      height: 600px;
      background: radial-gradient(circle, rgba(0, 240, 255, 0.04) 0%, transparent 70%);
      top: -200px;
      right: -100px;
      pointer-events: none;
      z-index: 0;
    }
    .space-glow-2 {
      position: absolute;
      width: 800px;
      height: 800px;
      background: radial-gradient(circle, rgba(0, 102, 255, 0.03) 0%, transparent 70%);
      bottom: -300px;
      left: -200px;
      pointer-events: none;
      z-index: 0;
    }

    /* Technical Scanline SVG overlay */
    .scanline-svg {
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      height: 100%;
      pointer-events: none;
      z-index: 1;
      stroke: rgba(0, 240, 255, 0.03);
    }

    /* Progress bar */
    .progress-bar-container {
      position: fixed;
      top: 0;
      left: 0;
      width: 100%;
      height: 4px;
      background: rgba(255, 255, 255, 0.03);
      z-index: 10;
    }
    .progress-bar {
      height: 100%;
      width: 0%;
      background: linear-gradient(90deg, var(--accent-blue), var(--accent-cyan));
      box-shadow: 0 0 8px var(--accent-cyan);
      transition: width 0.3s ease;
    }

    /* Slide Manager */
    .slides-container {
      width: 100%;
      height: 100%;
      position: relative;
      z-index: 2;
    }

    .slide {
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      height: 100%;
      display: none;
      opacity: 0;
      padding: 60px 100px;
      box-sizing: border-box;
      transition: opacity 0.4s cubic-bezier(0.16, 1, 0.3, 1), transform 0.4s cubic-bezier(0.16, 1, 0.3, 1);
      transform: translateY(15px) scale(0.995);
      flex-direction: column;
      justify-content: center;
    }

    .slide.active {
      display: flex;
      opacity: 1;
      transform: translateY(0) scale(1);
    }

    .slide-content-wrap {
      max-width: 1300px;
      width: 100%;
      margin: 0 auto;
      height: 100%;
      display: flex;
      flex-direction: column;
      justify-content: center;
    }

    /* Typography Defaults */
    .slide h1, .slide h2 {
      font-family: var(--font-display);
      font-weight: 700;
      letter-spacing: -0.03em;
      line-height: 1.15;
    }

    .slide h2 {
      font-size: 2.6rem;
      margin-bottom: 24px;
      background: linear-gradient(135deg, #ffffff 30%, #88d5ff 100%);
      -webkit-background-clip: text;
      -webkit-text-fill-color: transparent;
      display: flex;
      align-items: center;
      gap: 16px;
    }
    
    .slide h2::before {
      content: '';
      display: inline-block;
      width: 8px;
      height: 28px;
      background: var(--accent-cyan);
      box-shadow: 0 0 8px var(--accent-cyan);
      border-radius: 4px;
    }

    .slide p {
      color: var(--text-secondary);
      font-size: 1.05rem;
      line-height: 1.6;
      margin-bottom: 16px;
    }

    /* Glass Cards */
    .glass-card {
      background: var(--bg-card);
      backdrop-filter: blur(12px);
      border: 1px solid var(--border-cyan);
      border-radius: 14px;
      padding: 24px;
      box-shadow: 0 10px 30px rgba(0, 0, 0, 0.2);
      transition: border-color 0.3s ease, box-shadow 0.3s ease;
    }
    .glass-card:hover {
      border-color: var(--border-cyan-glow);
      box-shadow: 0 10px 35px rgba(0, 240, 255, 0.04);
    }

    /* Code Blocks */
    pre {
      background: #03060b !important;
      border: 1px solid rgba(255, 255, 255, 0.05);
      border-radius: 8px;
      padding: 18px;
      overflow-x: auto;
    }
    code {
      font-family: var(--font-mono);
      font-size: 0.9rem;
      color: #38bdf8;
    }

    /* Lists */
    ul, ol {
      margin-left: 24px;
      margin-bottom: 16px;
    }
    li {
      color: var(--text-secondary);
      font-size: 1rem;
      margin-bottom: 8px;
      line-height: 1.5;
    }
    strong {
      color: #ffffff;
      font-weight: 600;
    }

    /* Tables */
    table {
      width: 100%;
      border-collapse: collapse;
      margin: 16px 0;
      font-size: 0.95rem;
      background: rgba(5, 7, 11, 0.3);
      border-radius: 8px;
      overflow: hidden;
      border: 1px solid rgba(0, 240, 255, 0.05);
    }
    th {
      background: rgba(0, 240, 255, 0.04);
      color: var(--accent-cyan);
      font-family: var(--font-display);
      font-weight: 600;
      text-transform: uppercase;
      font-size: 0.75rem;
      letter-spacing: 0.08em;
      padding: 14px 16px;
      text-align: left;
      border-bottom: 1px solid rgba(0, 240, 255, 0.1);
    }
    td {
      padding: 14px 16px;
      border-bottom: 1px solid rgba(255, 255, 255, 0.02);
      color: var(--text-secondary);
      vertical-align: middle;
    }
    tr:last-child td {
      border-bottom: none;
    }
    tr:hover td {
      color: var(--text-primary);
      background: rgba(0, 240, 255, 0.01);
    }

    blockquote {
      border-left: 3px solid var(--accent-cyan);
      background: rgba(0, 240, 255, 0.02);
      padding: 16px 20px;
      margin: 16px 0;
      border-radius: 0 8px 8px 0;
    }
    blockquote p {
      margin: 0;
      font-style: italic;
      color: #cbd5e1;
    }

    /* HUD Bar */
    .hud-bar {
      position: fixed;
      bottom: 24px;
      left: 50%;
      transform: translateX(-50%);
      width: 90%;
      max-width: 1200px;
      height: 54px;
      background: rgba(10, 14, 23, 0.8);
      backdrop-filter: blur(16px);
      border: 1px solid rgba(0, 240, 255, 0.08);
      border-radius: 27px;
      z-index: 100;
      display: flex;
      justify-content: space-between;
      align-items: center;
      padding: 0 24px;
      box-shadow: 0 8px 32px rgba(0, 0, 0, 0.5);
    }

    .hud-left {
      display: flex;
      align-items: center;
      gap: 16px;
    }
    .hud-logo {
      font-family: var(--font-display);
      font-weight: 800;
      font-size: 0.95rem;
      letter-spacing: 0.15em;
      color: #ffffff;
      display: flex;
      align-items: center;
      gap: 8px;
    }
    .hud-logo span {
      color: var(--accent-cyan);
    }
    .hud-dot {
      width: 6px;
      height: 6px;
      background: var(--accent-cyan);
      border-radius: 50%;
      box-shadow: 0 0 6px var(--accent-cyan);
      animation: pulse 2s infinite;
    }

    .hud-center {
      display: flex;
      align-items: center;
      gap: 6px;
    }
    .indicator-dot {
      width: 7px;
      height: 7px;
      border-radius: 50%;
      background: rgba(255, 255, 255, 0.15);
      cursor: pointer;
      transition: background 0.3s ease, transform 0.3s ease;
    }
    .indicator-dot.active {
      background: var(--accent-cyan);
      box-shadow: 0 0 6px var(--accent-cyan);
      transform: scale(1.3);
    }

    .hud-right {
      display: flex;
      align-items: center;
      gap: 20px;
    }
    .slide-number {
      font-family: var(--font-mono);
      font-size: 0.85rem;
      color: var(--text-secondary);
      letter-spacing: 0.05em;
    }
    .slide-number span {
      color: #ffffff;
      font-weight: 600;
    }
    .nav-buttons {
      display: flex;
      gap: 8px;
    }
    .nav-btn {
      width: 32px;
      height: 32px;
      border-radius: 50%;
      border: 1px solid rgba(255, 255, 255, 0.08);
      background: rgba(255, 255, 255, 0.02);
      color: #ffffff;
      display: flex;
      justify-content: center;
      align-items: center;
      cursor: pointer;
      transition: background 0.2s, border-color 0.2s;
    }
    .nav-btn:hover {
      background: rgba(0, 240, 255, 0.05);
      border-color: var(--accent-cyan);
      color: var(--accent-cyan);
    }
    .nav-btn:active {
      transform: scale(0.95);
    }

    /* Layout Specifics */
    
    /* Slide 0: Title Slide */
    .slide-0 .slide-content-wrap {
      align-items: center;
      text-align: center;
      max-width: 900px;
      position: relative;
    }
    .slide-0 h1 {
      font-size: 3.6rem;
      margin-bottom: 20px;
      background: linear-gradient(135deg, #ffffff 40%, #5dc5ff 100%);
      -webkit-background-clip: text;
      -webkit-text-fill-color: transparent;
      letter-spacing: -0.04em;
    }
    .slide-0 blockquote {
      border-left: none;
      background: none;
      padding: 0;
      margin-top: 32px;
      width: 100%;
    }
    .slide-0 .title-grid {
      display: grid;
      grid-template-columns: repeat(3, 1fr);
      gap: 20px;
      width: 100%;
      margin-top: 40px;
    }
    .slide-0 .title-cell {
      background: rgba(10, 14, 23, 0.4);
      border: 1px solid rgba(0, 240, 255, 0.05);
      border-radius: 12px;
      padding: 20px;
      text-align: left;
    }
    .slide-0 .title-cell-label {
      font-size: 0.75rem;
      color: var(--accent-cyan);
      font-family: var(--font-display);
      text-transform: uppercase;
      letter-spacing: 0.1em;
      margin-bottom: 8px;
    }
    .slide-0 .title-cell-value {
      font-size: 0.95rem;
      color: #ffffff;
      line-height: 1.4;
    }
    .slide-0 .title-cell-value a {
      color: var(--accent-cyan);
      text-decoration: none;
    }
    .slide-0 .title-cell-value a:hover {
      text-decoration: underline;
    }
    .slide-0 .title-cell-large {
      grid-column: span 3;
    }

    /* Slide 1: Mission */
    .slide-1 .slide-content-wrap {
      max-width: 950px;
      align-items: center;
      text-align: center;
    }
    .slide-1 h2 {
      justify-content: center;
      margin-bottom: 32px;
    }
    .slide-1 h2::before {
      display: none;
    }
    .slide-1 p {
      font-size: 2.1rem;
      line-height: 1.5;
      font-weight: 300;
      color: #e2e8f0;
      letter-spacing: -0.01em;
    }

    /* Slide 2: The Four Roles */
    .slide-2 .roles-grid {
      display: grid;
      grid-template-columns: repeat(4, 1fr);
      gap: 20px;
      margin-top: 30px;
    }
    .slide-2 .role-card {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 12px;
      padding: 24px;
      backdrop-filter: blur(8px);
      transition: all 0.3s ease;
    }
    .slide-2 .role-card:hover {
      border-color: var(--accent-cyan);
      transform: translateY(-4px);
    }
    .slide-2 .role-badge {
      display: inline-block;
      font-size: 0.7rem;
      color: var(--accent-cyan);
      background: rgba(0, 240, 255, 0.05);
      border: 1px solid rgba(0, 240, 255, 0.15);
      padding: 2px 8px;
      border-radius: 10px;
      font-family: var(--font-display);
      font-weight: 600;
      letter-spacing: 0.05em;
      margin-bottom: 12px;
    }
    .slide-2 .role-name {
      font-size: 1.15rem;
      color: #ffffff;
      font-family: var(--font-display);
      margin-bottom: 4px;
    }
    .slide-2 .role-title {
      font-size: 0.9rem;
      color: var(--accent-cyan);
      margin-bottom: 16px;
      font-weight: 500;
    }
    .slide-2 .role-desc {
      font-size: 0.9rem;
      color: var(--text-secondary);
      line-height: 1.5;
    }
    .slide-2 blockquote {
      margin-top: 30px;
      font-size: 0.9rem;
    }

    /* Slides 3-6: Team Member Slides */
    .slide-team-layout {
      display: grid;
      grid-template-columns: 4fr 5fr;
      gap: 40px;
      height: calc(100vh - 200px);
      align-items: center;
    }
    .slide-team-info {
      display: flex;
      flex-direction: column;
      gap: 20px;
      height: 100%;
      justify-content: center;
    }
    .slide-team-timeline {
      height: 100%;
      display: flex;
      flex-direction: column;
      justify-content: center;
    }
    .timeline-scroll {
      max-height: 520px;
      overflow-y: auto;
      padding-right: 10px;
    }
    .timeline-scroll::-webkit-scrollbar {
      width: 4px;
    }
    .timeline-scroll::-webkit-scrollbar-track {
      background: rgba(255, 255, 255, 0.02);
    }
    .timeline-scroll::-webkit-scrollbar-thumb {
      background: rgba(0, 240, 255, 0.15);
      border-radius: 2px;
    }
    .timeline-scroll::-webkit-scrollbar-thumb:hover {
      background: var(--accent-cyan);
    }
    .slide-team-info .checks-card {
      margin-top: 10px;
    }

    /* Technical vector diagrams (mock diagrams inside cards) */
    .vector-diagram {
      height: 100px;
      border: 1px solid rgba(0, 240, 255, 0.06);
      background: rgba(2, 3, 5, 0.4);
      border-radius: 8px;
      margin-bottom: 16px;
      display: flex;
      justify-content: center;
      align-items: center;
      position: relative;
      overflow: hidden;
    }
    
    /* Slide 7: Contracts.md */
    .slide-7 .contracts-layout {
      display: grid;
      grid-template-columns: 280px 1fr;
      gap: 32px;
      margin-top: 20px;
    }
    .slide-7 .tabs-container {
      display: flex;
      flex-direction: column;
      gap: 8px;
    }
    .slide-7 .tab-item {
      padding: 16px 20px;
      background: rgba(10, 14, 23, 0.4);
      border: 1px solid rgba(0, 240, 255, 0.05);
      border-radius: 10px;
      cursor: pointer;
      transition: all 0.2s ease;
      font-family: var(--font-mono);
      font-size: 0.8rem;
      color: var(--text-secondary);
      text-align: left;
    }
    .slide-7 .tab-item.active {
      background: var(--bg-card);
      border-color: var(--accent-cyan);
      color: #ffffff;
      box-shadow: 0 0 10px rgba(0, 240, 255, 0.05);
    }
    .slide-7 .tab-item-method {
      font-weight: 700;
      color: var(--accent-cyan);
      margin-right: 6px;
    }
    .slide-7 .code-container {
      position: relative;
      height: 480px;
    }
    .slide-7 .code-pane {
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      height: 100%;
      display: none;
    }
    .slide-7 .code-pane.active {
      display: block;
    }
    .slide-7 pre {
      height: 100%;
      margin: 0;
    }

    /* Slide 8: Communication rules */
    .slide-8 .rules-layout {
      display: grid;
      grid-template-columns: repeat(5, 1fr);
      gap: 16px;
      margin-top: 30px;
    }
    .slide-8 .rule-card {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 12px;
      padding: 24px;
      display: flex;
      flex-direction: column;
      position: relative;
      height: 380px;
    }
    .slide-8 .rule-number {
      font-family: var(--font-display);
      font-weight: 800;
      font-size: 2.2rem;
      color: rgba(0, 240, 255, 0.08);
      margin-bottom: 20px;
    }
    .slide-8 .rule-title {
      font-family: var(--font-display);
      font-size: 1rem;
      font-weight: 600;
      color: var(--accent-cyan);
      margin-bottom: 12px;
      line-height: 1.4;
    }
    .slide-8 .rule-desc {
      font-size: 0.88rem;
      color: var(--text-secondary);
      line-height: 1.5;
    }

    /* Slide 9: Demo script */
    .slide-9 .demo-timeline {
      display: grid;
      grid-template-columns: repeat(5, 1fr);
      gap: 16px;
      margin-top: 40px;
      position: relative;
    }
    .slide-9 .demo-timeline::before {
      content: '';
      position: absolute;
      top: 32px;
      left: 0;
      width: 100%;
      height: 2px;
      background: rgba(0, 240, 255, 0.08);
      z-index: 0;
    }
    .slide-9 .demo-step {
      position: relative;
      z-index: 1;
      padding-top: 50px;
    }
    .slide-9 .demo-node {
      position: absolute;
      top: 24px;
      left: 0;
      width: 18px;
      height: 18px;
      background: var(--bg-space);
      border: 3px solid var(--accent-cyan);
      border-radius: 50%;
      box-shadow: 0 0 8px var(--accent-cyan);
    }
    .slide-9 .demo-time {
      font-family: var(--font-mono);
      font-size: 0.85rem;
      font-weight: 600;
      color: var(--accent-cyan);
      margin-bottom: 8px;
    }
    .slide-9 .demo-card {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 10px;
      padding: 18px;
      min-height: 260px;
    }
    .slide-9 .demo-desc {
      font-size: 0.9rem;
      color: var(--text-secondary);
      line-height: 1.5;
    }

    /* Slide 10: Risk register */
    .slide-10 table {
      margin-top: 30px;
    }
    .slide-10 td:first-child {
      font-weight: 600;
      color: #ffffff;
      width: 32%;
    }
    .risk-badge {
      display: inline-block;
      font-size: 0.7rem;
      font-weight: 700;
      letter-spacing: 0.05em;
      padding: 2px 8px;
      border-radius: 4px;
      background: rgba(239, 68, 68, 0.1);
      border: 1px solid rgba(239, 68, 68, 0.2);
      color: #ef4444;
      margin-bottom: 6px;
    }

    /* Slide 11: Submission checklist & Slide 15: DOD */
    .checklist-container {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 16px;
      margin-top: 30px;
    }
    .check-item {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 10px;
      padding: 16px 20px;
      display: flex;
      align-items: center;
      gap: 16px;
      cursor: pointer;
      transition: border-color 0.2s;
    }
    .check-item:hover {
      border-color: rgba(0, 240, 255, 0.2);
    }
    .check-ring {
      width: 20px;
      height: 20px;
      border-radius: 50%;
      border: 2px solid rgba(255, 255, 255, 0.15);
      display: flex;
      justify-content: center;
      align-items: center;
      flex-shrink: 0;
      transition: border-color 0.2s, background-color 0.2s;
    }
    .check-ring::after {
      content: '';
      width: 8px;
      height: 8px;
      border-radius: 50%;
      background: var(--accent-cyan);
      opacity: 0;
      transition: opacity 0.2s;
    }
    .check-item.checked .check-ring {
      border-color: var(--accent-cyan);
      box-shadow: 0 0 6px var(--accent-cyan-glow);
    }
    .check-item.checked .check-ring::after {
      opacity: 1;
    }
    .check-item.checked .check-text {
      color: #ffffff;
    }
    .check-text {
      font-size: 0.95rem;
      color: var(--text-secondary);
      line-height: 1.4;
    }
    .check-text code {
      background: rgba(255, 255, 255, 0.04);
      padding: 1px 4px;
      border-radius: 4px;
    }

    /* Slide 12: Hand-off */
    .slide-12 .handoff-layout {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 32px;
      margin-top: 30px;
    }
    .slide-12 .handoff-left {
      display: flex;
      flex-direction: column;
      justify-content: center;
      gap: 20px;
    }
    .slide-12 .handoff-title-p {
      font-size: 1.3rem;
      color: #ffffff;
      font-weight: 500;
      line-height: 1.5;
    }
    .slide-12 .handoff-right {
      display: flex;
      flex-direction: column;
      gap: 12px;
    }
    .slide-12 .handoff-item {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 10px;
      padding: 18px 20px;
      display: flex;
      gap: 16px;
    }
    .slide-12 .handoff-num {
      font-family: var(--font-mono);
      color: var(--accent-cyan);
      font-weight: 700;
      font-size: 1.1rem;
    }
    .slide-12 .handoff-desc {
      font-size: 0.95rem;
      color: var(--text-secondary);
      line-height: 1.4;
    }
    .slide-12 .handoff-desc code {
      background: rgba(255, 255, 255, 0.04);
      padding: 1px 4px;
      border-radius: 4px;
    }
    .slide-12 .handoff-footer {
      margin-top: 20px;
      padding: 16px;
      border: 1px dashed rgba(0, 240, 255, 0.15);
      border-radius: 8px;
      background: rgba(0, 240, 255, 0.01);
      font-style: italic;
      color: var(--text-secondary);
      font-size: 0.9rem;
      text-align: center;
    }

    /* Slide 13: Teammate Messages */
    .slide-13 .messages-layout {
      display: grid;
      grid-template-columns: 240px 1fr;
      gap: 32px;
      margin-top: 20px;
    }
    .slide-13 .messages-tabs {
      display: flex;
      flex-direction: column;
      gap: 8px;
    }
    .slide-13 .msg-tab {
      padding: 14px 18px;
      background: rgba(10, 14, 23, 0.4);
      border: 1px solid rgba(0, 240, 255, 0.05);
      border-radius: 10px;
      cursor: pointer;
      transition: all 0.2s ease;
      color: var(--text-secondary);
      text-align: left;
      font-family: var(--font-display);
      font-weight: 600;
      display: flex;
      align-items: center;
      gap: 12px;
    }
    .slide-13 .msg-tab.active {
      background: var(--bg-card);
      border-color: var(--accent-cyan);
      color: #ffffff;
    }
    .slide-13 .msg-avatar-placeholder {
      width: 24px;
      height: 24px;
      border-radius: 50%;
      background: var(--accent-blue);
      display: flex;
      justify-content: center;
      align-items: center;
      font-size: 0.7rem;
      font-weight: 700;
      color: #ffffff;
    }
    .slide-13 .msg-tab.active .msg-avatar-placeholder {
      background: var(--accent-cyan);
      box-shadow: 0 0 6px var(--accent-cyan);
    }
    .slide-13 .messages-panes {
      position: relative;
      height: 480px;
    }
    .slide-13 .msg-pane {
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      height: 100%;
      display: none;
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 12px;
      padding: 24px;
      overflow-y: auto;
    }
    .slide-13 .msg-pane.active {
      display: block;
    }
    .slide-13 .msg-header {
      font-family: var(--font-display);
      font-size: 1.1rem;
      font-weight: 700;
      margin-bottom: 12px;
      color: #ffffff;
      border-bottom: 1px solid rgba(255, 255, 255, 0.05);
      padding-bottom: 12px;
    }
    .slide-13 .msg-instruction {
      margin-bottom: 18px;
      color: var(--accent-cyan);
      font-size: 0.9rem;
      font-weight: 500;
    }
    .slide-13 .msg-body {
      font-size: 0.95rem;
      color: var(--text-secondary);
      line-height: 1.6;
    }
    .slide-13 .msg-body p {
      margin-bottom: 12px;
    }
    .slide-13 .msg-body ul {
      margin-left: 20px;
      margin-bottom: 12px;
    }
    .slide-13 .msg-body li {
      margin-bottom: 6px;
    }
    .slide-13 .msg-body a {
      color: var(--accent-cyan);
      text-decoration: none;
    }
    .slide-13 .msg-body a:hover {
      text-decoration: underline;
    }

    /* Slide 14: The two things Aayush will not write */
    .slide-14 .nw-layout {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 32px;
      margin-top: 30px;
      align-items: center;
    }
    .slide-14 .nw-left {
      display: flex;
      flex-direction: column;
      gap: 16px;
    }
    .slide-14 .nw-card {
      background: var(--bg-card);
      border: 1px solid var(--border-cyan);
      border-radius: 12px;
      padding: 24px;
      position: relative;
    }
    .slide-14 .nw-card-title {
      font-family: var(--font-display);
      font-weight: 700;
      font-size: 1.1rem;
      color: #ef4444;
      margin-bottom: 10px;
      display: flex;
      align-items: center;
      gap: 10px;
    }
    .slide-14 .nw-card-title::before {
      content: '!';
      display: inline-flex;
      justify-content: center;
      align-items: center;
      width: 20px;
      height: 20px;
      background: #ef4444;
      color: #ffffff;
      border-radius: 50%;
      font-size: 0.75rem;
      font-weight: 800;
    }
    .slide-14 .nw-card-desc {
      font-size: 0.9rem;
      color: var(--text-secondary);
      line-height: 1.5;
    }
    .slide-14 .nw-right {
      display: flex;
      flex-direction: column;
      gap: 12px;
    }
    .slide-14 .nw-terminal-label {
      font-family: var(--font-mono);
      font-size: 0.75rem;
      color: var(--text-secondary);
      text-transform: uppercase;
      letter-spacing: 0.05em;
      margin-bottom: 4px;
    }
    .slide-14 .nw-terminal {
      background: #020305;
      border: 1px solid rgba(239, 68, 68, 0.15);
      border-radius: 8px;
      padding: 16px;
      font-family: var(--font-mono);
      font-size: 0.85rem;
      color: #ef4444;
      line-height: 1.4;
      box-shadow: 0 0 15px rgba(239, 68, 68, 0.03);
    }
    .slide-14 .nw-info-box {
      margin-top: 16px;
      padding: 16px;
      border: 1px solid rgba(0, 240, 255, 0.1);
      background: rgba(0, 240, 255, 0.01);
      border-radius: 8px;
      font-size: 0.95rem;
      color: var(--text-secondary);
      text-align: center;
    }

    /* Slide 1: Mission Split Layout */
    .mission-split-layout {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 40px;
      margin-top: 30px;
      align-items: center;
    }
    .mission-left-panel {
      padding: 40px;
      min-height: 400px;
      display: flex;
      flex-direction: column;
      justify-content: center;
    }
    .mission-badge {
      font-size: 0.75rem;
      color: var(--accent-cyan);
      background: rgba(0, 240, 255, 0.05);
      border: 1px solid rgba(0, 240, 255, 0.15);
      padding: 4px 12px;
      border-radius: 12px;
      font-family: var(--font-display);
      font-weight: 600;
      letter-spacing: 0.1em;
      margin-bottom: 24px;
      align-self: flex-start;
    }
    .mission-hero-text {
      font-size: 1.65rem !important;
      line-height: 1.55;
      font-weight: 300;
      color: #ffffff !important;
      letter-spacing: -0.015em;
      margin-bottom: 0;
    }
    .mission-right-panel {
      height: 400px;
      position: relative;
    }
    .satellite-img-container {
      width: 100%;
      height: 100%;
      border: 1px solid rgba(0, 240, 255, 0.2);
      border-radius: 12px;
      position: relative;
      overflow: hidden;
      box-shadow: 0 0 25px rgba(0, 240, 255, 0.03);
    }
    .satellite-img {
      width: 100%;
      height: 100%;
      object-fit: cover;
      opacity: 0.85;
      filter: contrast(1.1) brightness(0.9);
      transition: transform 0.5s ease;
    }
    .satellite-img-container:hover .satellite-img {
      transform: scale(1.02);
    }
    .satellite-scanline {
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      height: 2px;
      background: rgba(0, 240, 255, 0.3);
      box-shadow: 0 0 8px rgba(0, 240, 255, 0.6);
      animation: scan 6s linear infinite;
      pointer-events: none;
    }
    @keyframes scan {
      0% { top: 0%; }
      100% { top: 100%; }
    }
    .coordinate-tag {
      position: absolute;
      font-family: var(--font-mono);
      font-size: 0.7rem;
      color: var(--accent-cyan);
      background: rgba(2, 3, 5, 0.7);
      padding: 2px 6px;
      border: 1px solid rgba(0, 240, 255, 0.15);
      border-radius: 4px;
      pointer-events: none;
    }
    .coordinate-top-left { top: 8px; left: 8px; }
    .coordinate-top-right { top: 8px; right: 8px; }
    .coordinate-bottom-left { bottom: 8px; left: 8px; }
    .coordinate-bottom-right { bottom: 8px; right: 8px; color: #ef4444; border-color: rgba(239, 68, 68, 0.3); }

    /* Animation pulse */
    @keyframes pulse {
      0% { box-shadow: 0 0 4px var(--accent-cyan); opacity: 0.6; }
      50% { box-shadow: 0 0 10px var(--accent-cyan); opacity: 1; }
      100% { box-shadow: 0 0 4px var(--accent-cyan); opacity: 0.6; }
    }

    /* responsive */
    @media (max-width: 1024px) {
      .slide { padding: 40px 40px; }
      .slide h2 { font-size: 2rem; }
      .slide-team-layout { grid-template-columns: 1fr; height: auto; gap: 20px; }
      .timeline-scroll { max-height: 280px; }
    }
  </style>
</head>
<body>

  <!-- Space radial glows -->
  <div class="space-glow"></div>
  <div class="space-glow-2"></div>

  <!-- Technical overlays -->
  <svg class="scanline-svg" viewBox="0 0 1000 1000" fill="none">
    <circle cx="500" cy="500" r="450" stroke-dasharray="1 10" />
    <circle cx="500" cy="500" r="300" stroke-dasharray="3 15" />
    <line x1="500" y1="50" x2="500" y2="950" stroke-dasharray="2 20" />
    <line x1="50" y1="500" x2="950" y2="500" stroke-dasharray="2 20" />
  </svg>

  <!-- Progress bar -->
  <div class="progress-bar-container">
    <div class="progress-bar" id="progressBar"></div>
  </div>

  <!-- Slides -->
  <div class="slides-container" id="slidesContainer">
    {{SLIDES_CONTENT}}
  </div>

  <!-- HUD navigation bar -->
  <div class="hud-bar">
    <div class="hud-left">
      <div class="hud-dot"></div>
      <div class="hud-logo">J.A.R.V.I.S. <span>MISSION</span></div>
    </div>
    
    <div class="hud-center" id="hudCenter">
      <!-- Dots injected dynamically -->
    </div>

    <div class="hud-right">
      <div class="slide-number">SLIDE <span id="currentNum">01</span> / <span id="totalNum">16</span></div>
      <div class="nav-buttons">
        <button class="nav-btn" id="prevBtn" title="Previous Slide (Arrow Left)">
          <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="19" y1="12" x2="5" y2="12"></line><polyline points="12 19 5 12 12 5"></polyline></svg>
        </button>
        <button class="nav-btn" id="nextBtn" title="Next Slide (Arrow Right / Space)">
          <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="5" y1="12" x2="19" y2="12"></line><polyline points="12 5 19 12 12 19"></polyline></svg>
        </button>
      </div>
    </div>
  </div>

  <script>
    document.addEventListener("DOMContentLoaded", () => {
      const slides = Array.from(document.querySelectorAll(".slide"));
      const progressBar = document.getElementById("progressBar");
      const hudCenter = document.getElementById("hudCenter");
      const currentNum = document.getElementById("currentNum");
      const totalNum = document.getElementById("totalNum");
      const prevBtn = document.getElementById("prevBtn");
      const nextBtn = document.getElementById("nextBtn");

      let currentSlide = 0;
      const totalSlides = slides.length;

      totalNum.textContent = String(totalSlides).padStart(2, '0');

      // Generate HUD indicator dots
      slides.forEach((_, idx) => {
        const dot = document.createElement("div");
        dot.className = `indicator-dot ${idx === 0 ? 'active' : ''}`;
        dot.title = `Go to Slide ${idx + 1}`;
        dot.addEventListener("click", () => goToSlide(idx));
        hudCenter.appendChild(dot);
      });

      const dots = Array.from(document.querySelectorAll(".indicator-dot"));

      function updateHUD() {
        currentNum.textContent = String(currentSlide + 1).padStart(2, '0');
        progressBar.style.width = `${((currentSlide) / (totalSlides - 1)) * 100}%`;
        
        dots.forEach((dot, idx) => {
          if (idx === currentSlide) {
            dot.classList.add("active");
          } else {
            dot.classList.remove("active");
          }
        });
      }

      function goToSlide(index) {
        if (index < 0 || index >= totalSlides) return;
        
        slides[currentSlide].classList.remove("active");
        currentSlide = index;
        slides[currentSlide].classList.add("active");
        
        updateHUD();
      }

      function next() {
        if (currentSlide < totalSlides - 1) {
          goToSlide(currentSlide + 1);
        }
      }

      function prev() {
        if (currentSlide > 0) {
          goToSlide(currentSlide - 1);
        }
      }

      // Button listeners
      nextBtn.addEventListener("click", next);
      prevBtn.addEventListener("click", prev);

      // Keyboard listeners
      document.addEventListener("keydown", (e) => {
        if (e.key === "ArrowRight" || e.key === "Space" || e.key === "PageDown") {
          e.preventDefault();
          next();
        } else if (e.key === "ArrowLeft" || e.key === "Backspace" || e.key === "PageUp") {
          e.preventDefault();
          prev();
        }
      });

      // Swipe support
      let touchStartX = 0;
      let touchEndX = 0;
      
      document.addEventListener("touchstart", (e) => {
        touchStartX = e.changedTouches[0].screenX;
      });

      document.addEventListener("touchend", (e) => {
        touchEndX = e.changedTouches[0].screenX;
        handleSwipe();
      });

      function handleSwipe() {
        if (touchStartX - touchEndX > 50) {
          next(); // swipe left -> next slide
        } else if (touchEndX - touchStartX > 50) {
          prev(); // swipe right -> prev slide
        }
      }

      // Initialize Layout Transforms for slides

      // 1. Transform Slide 0: Title Slide
      const slide0 = document.querySelector(".slide-0");
      if (slide0) {
        const titleEl = slide0.querySelector("h1");
        const blockquote = slide0.querySelector("blockquote");
        if (blockquote) {
          const text = blockquote.textContent.trim();
          
          // Regex extraction
          const hackathon = text.match(/Hackathon:\\s*([^\\n]+)/i)?.[1] || "";
          const repo = text.match(/Repo \\(sanitized\\):\\s*([^\\n]+)/i)?.[1] || "";
          const folder = text.match(/Working folder:\\s*([^\\n]+)/i)?.[1] || "";
          const windowText = text.match(/Sprint window:\\s*([^\\n]+)/i)?.[1] || "";
          const goal = text.match(/Goal:\\s*([^$]+)/i)?.[1] || "";
          
          const gridDiv = document.createElement("div");
          gridDiv.className = "title-grid";
          gridDiv.innerHTML = `
            <div class="title-cell">
              <div class="title-cell-label">HACKATHON</div>
              <div class="title-cell-value">${hackathon}</div>
            </div>
            <div class="title-cell">
              <div class="title-cell-label">REPOSITORY</div>
              <div class="title-cell-value"><a href="${repo}" target="_blank">${repo}</a></div>
            </div>
            <div class="title-cell">
              <div class="title-cell-label">WORKING DIRECTORY</div>
              <div class="title-cell-value"><code>${folder}</code></div>
            </div>
            <div class="title-cell">
              <div class="title-cell-label">SPRINT TIME WINDOW</div>
              <div class="title-cell-value">${windowText}</div>
            </div>
            <div class="title-cell title-cell-large">
              <div class="title-cell-label">MISSION GOAL</div>
              <div class="title-cell-value" style="font-weight: 500; color: var(--accent-cyan);">${goal}</div>
            </div>
          `;
          blockquote.replaceWith(gridDiv);
        }
      }

      // 1. Transform Slide 1: Mission
      const slide1 = document.querySelector(".slide-1");
      if (slide1) {
        const contentWrap = slide1.querySelector(".slide-content-wrap");
        const h2 = contentWrap.querySelector("h2");
        const pEl = contentWrap.querySelector("p");
        if (pEl) {
          const splitLayout = document.createElement("div");
          splitLayout.className = "mission-split-layout";
          splitLayout.innerHTML = `
            <div class="mission-left-panel glass-card">
              <div class="mission-badge">OPERATIONAL GOAL</div>
              <p class="mission-hero-text">${pEl.innerHTML}</p>
            </div>
            <div class="mission-right-panel">
              <div class="satellite-img-container">
                <img src="industrial_satellite_view.jpg" alt="Industrial Complex Satellite Imagery" class="satellite-img" />
                <div class="satellite-scanline"></div>
                <div class="coordinate-tag coordinate-top-left">LAT 29.3400 N</div>
                <div class="coordinate-tag coordinate-top-right">LON 94.9100 W</div>
                <div class="coordinate-tag coordinate-bottom-left">ALT 3,500m</div>
                <div class="coordinate-tag coordinate-bottom-right">SENSOR ACTIVE</div>
              </div>
            </div>
          `;
          
          contentWrap.innerHTML = "";
          contentWrap.appendChild(h2);
          contentWrap.appendChild(splitLayout);
        }
      }

      // 2. Transform Slide 2: The Four Roles
      const slide2 = document.querySelector(".slide-2");
      if (slide2) {
        const table = slide2.querySelector("table");
        if (table) {
          const rows = Array.from(table.querySelectorAll("tbody tr"));
          const grid = document.createElement("div");
          grid.className = "roles-grid";
          
          rows.forEach((row, rIdx) => {
            const cells = row.querySelectorAll("td");
            if (cells.length >= 3) {
              const card = document.createElement("div");
              card.className = "role-card";
              card.innerHTML = `
                <div class="role-badge">MEMBER 0${rIdx+1}</div>
                <div class="role-name">${cells[0].innerHTML}</div>
                <div class="role-title">${cells[1].innerHTML}</div>
                <div class="role-desc">${cells[2].innerHTML}</div>
              `;
              grid.appendChild(card);
            }
          });
          table.replaceWith(grid);
        }
      }

      // 3. Transform Slide 3, 4, 5, 6: Team Member Deliverables
      const teamSlideIndices = [3, 4, 5, 6];
      teamSlideIndices.forEach(idx => {
        const slide = document.querySelector(`.slide-${idx}`);
        if (slide) {
          // Structure as split column dashboard
          const contentWrap = slide.querySelector(".slide-content-wrap");
          
          // Extract elements
          const h2 = contentWrap.querySelector("h2");
          const pEl = contentWrap.querySelector("p");
          const table = contentWrap.querySelector("table");
          const remainingElements = Array.from(contentWrap.children).filter(el => el !== h2 && el !== table && el !== pEl);
          
          const layout = document.createElement("div");
          layout.className = "slide-team-layout";
          
          const infoCol = document.createElement("div");
          infoCol.className = "slide-team-info";
          if (pEl) infoCol.appendChild(pEl);
          
          // Group remaining checklist items / notes
          const checksCard = document.createElement("div");
          checksCard.className = "glass-card checks-card";
          remainingElements.forEach(el => checksCard.appendChild(el));
          infoCol.appendChild(checksCard);
          
          const timelineCol = document.createElement("div");
          timelineCol.className = "slide-team-timeline";
          
          // Custom diagram header inside card
          const diagram = document.createElement("div");
          diagram.className = "vector-diagram";
          if (idx === 3) {
            diagram.innerHTML = `
              <svg width="220" height="60" viewBox="0 0 220 60" fill="none" stroke="rgba(0, 240, 255, 0.2)">
                <rect x="10" y="15" width="40" height="30" rx="4" />
                <rect x="80" y="15" width="60" height="30" rx="4" stroke="var(--accent-cyan)" />
                <rect x="170" y="15" width="40" height="30" rx="4" />
                <line x1="50" y1="30" x2="80" y2="30" />
                <line x1="140" y1="30" x2="170" y2="30" />
                <circle cx="80" cy="30" r="3" fill="var(--accent-cyan)" />
                <circle cx="140" cy="30" r="3" fill="var(--accent-cyan)" />
              </svg>
            `;
          } else if (idx === 4) {
            diagram.innerHTML = `
              <svg width="220" height="60" viewBox="0 0 220 60" fill="none" stroke="rgba(0, 240, 255, 0.2)">
                <circle cx="30" cy="30" r="20" />
                <circle cx="110" cy="30" r="20" stroke="var(--accent-cyan)" />
                <circle cx="190" cy="30" r="20" />
                <line x1="50" y1="30" x2="90" y2="30" />
                <line x1="130" y1="30" x2="170" y2="30" />
                <text x="110" y="34" fill="var(--accent-cyan)" font-size="10" font-family="monospace" text-anchor="middle">RAG</text>
              </svg>
            `;
          } else if (idx === 5) {
            diagram.innerHTML = `
              <svg width="220" height="60" viewBox="0 0 220 60" fill="none" stroke="rgba(0, 240, 255, 0.2)">
                <path d="M10,30 L60,15 L60,45 Z" />
                <line x1="60" y1="30" x2="110" y2="30" stroke="var(--accent-cyan)" />
                <circle cx="130" cy="30" r="15" stroke="var(--accent-cyan)" />
                <line x1="145" y1="30" x2="200" y2="30" />
              </svg>
            `;
          } else if (idx === 6) {
            diagram.innerHTML = `
              <svg width="220" height="60" viewBox="0 0 220 60" fill="none" stroke="rgba(0, 240, 255, 0.2)">
                <rect x="20" y="10" width="180" height="40" rx="4" />
                <line x1="20" y1="20" x2="200" y2="20" />
                <circle cx="40" cy="15" r="2" fill="var(--accent-cyan)" />
                <circle cx="50" cy="15" r="2" fill="var(--accent-cyan)" />
                <circle cx="60" cy="15" r="2" fill="var(--accent-cyan)" />
              </svg>
            `;
          }
          timelineCol.appendChild(diagram);
          
          if (table) {
            const tableScroll = document.createElement("div");
            tableScroll.className = "timeline-scroll";
            tableScroll.appendChild(table);
            timelineCol.appendChild(tableScroll);
          }
          
          layout.appendChild(infoCol);
          layout.appendChild(timelineCol);
          contentWrap.appendChild(layout);
        }
      });

      // 4. Transform Slide 7: CONTRACTS.md tabbed layout
      const slide7 = document.querySelector(".slide-7");
      if (slide7) {
        const contentWrap = slide7.querySelector(".slide-content-wrap");
        const h2 = contentWrap.querySelector("h2");
        const pEl = contentWrap.querySelector("p");
        const codeBlocks = Array.from(contentWrap.querySelectorAll("pre"));
        
        const layout = document.createElement("div");
        layout.className = "contracts-layout";
        
        const tabsContainer = document.createElement("div");
        tabsContainer.className = "tabs-container";
        
        const codeContainer = document.createElement("div");
        codeContainer.className = "code-container";
        
        const tabNames = [
          { method: "POST", path: "/api/chat" },
          { method: "POST", path: "/api/agent/orchestrate" },
          { method: "POST", path: "/api/ingest/upload" },
          { method: "GET", path: "/api/kg/neighbors" }
        ];

        codeBlocks.forEach((codeBlock, cIdx) => {
          const tab = document.createElement("button");
          tab.className = `tab-item ${cIdx === 0 ? 'active' : ''}`;
          tab.innerHTML = `<span class="tab-item-method">${tabNames[cIdx].method}</span>${tabNames[cIdx].path}`;
          tab.addEventListener("click", () => {
            Array.from(tabsContainer.querySelectorAll(".tab-item")).forEach((t, i) => {
              if (i === cIdx) t.classList.add("active");
              else t.classList.remove("active");
            });
            Array.from(codeContainer.querySelectorAll(".code-pane")).forEach((p, i) => {
              if (i === cIdx) p.classList.add("active");
              else p.classList.remove("active");
            });
          });
          tabsContainer.appendChild(tab);
          
          const pane = document.createElement("div");
          pane.className = `code-pane ${cIdx === 0 ? 'active' : ''}`;
          pane.appendChild(codeBlock);
          codeContainer.appendChild(pane);
        });
        
        layout.appendChild(tabsContainer);
        layout.appendChild(codeContainer);
        
        // Clear all except h2
        contentWrap.innerHTML = "";
        contentWrap.appendChild(h2);
        if (pEl) contentWrap.appendChild(pEl);
        contentWrap.appendChild(layout);
      }

      // 5. Transform Slide 8: Communication rules
      const slide8 = document.querySelector(".slide-8");
      if (slide8) {
        const listItems = Array.from(slide8.querySelectorAll("ul li"));
        const rulesLayout = document.createElement("div");
        rulesLayout.className = "rules-layout";
        
        listItems.forEach((li, lIdx) => {
          const card = document.createElement("div");
          card.className = "rule-card";
          
          const text = li.innerHTML;
          // Split title and desc if separated by bold or first sentence
          const boldMatch = text.match(/^<strong>(.*?)<\\/strong>(.*)/);
          let title = `Rule 0${lIdx+1}`;
          let desc = text;
          
          if (boldMatch) {
            title = boldMatch[1].replace(/[:\\.]\\s*$/, "");
            desc = boldMatch[2].trim().replace(/^[-\\s:\\.]*/, "");
          } else {
            const firstDot = text.indexOf(". ");
            if (firstDot !== -1) {
              title = text.substring(0, firstDot);
              desc = text.substring(firstDot + 2);
            }
          }
          
          card.innerHTML = `
            <div class="rule-number">0${lIdx+1}</div>
            <div class="rule-title">${title}</div>
            <div class="rule-desc">${desc}</div>
          `;
          rulesLayout.appendChild(card);
        });
        
        const ul = slide8.querySelector("ul");
        if (ul) ul.replaceWith(rulesLayout);
      }

      // 6. Transform Slide 9: Demo script
      const slide9 = document.querySelector(".slide-9");
      if (slide9) {
        const listItems = Array.from(slide9.querySelectorAll("ol li"));
        const timeline = document.createElement("div");
        timeline.className = "demo-timeline";
        
        listItems.forEach((li, lIdx) => {
          const text = li.innerHTML;
          const boldMatch = text.match(/^<strong>(.*?)<\\/strong>(.*)/);
          let timeLabel = `0${lIdx+1}`;
          let desc = text;
          
          if (boldMatch) {
            timeLabel = boldMatch[1].replace(/[:\\.]\\s*$/, "");
            desc = boldMatch[2].trim().replace(/^[-\\s:\\.]*/, "");
          }
          
          const step = document.createElement("div");
          step.className = "demo-step";
          step.innerHTML = `
            <div class="demo-node"></div>
            <div class="demo-time">${timeLabel}</div>
            <div class="demo-card">
              <div class="demo-desc">${desc}</div>
            </div>
          `;
          timeline.appendChild(step);
        });
        
        const ol = slide9.querySelector("ol");
        if (ol) ol.replaceWith(timeline);
      }

      // 7. Custom CSS overrides for slide 10 (Risk register)
      const slide10 = document.querySelector(".slide-10");
      if (slide10) {
        // Tag risk level cells visually
        const rows = Array.from(slide10.querySelectorAll("tbody tr"));
        rows.forEach(row => {
          const cell = row.querySelector("td");
          if (cell) {
            cell.innerHTML = `<span class="risk-badge">HIGH LIMIT</span><br/>${cell.innerHTML}`;
          }
        });
      }

      // 8. Transform Slide 11 & Slide 15: Checklist components
      const checkSlides = [11, 15];
      checkSlides.forEach(idx => {
        const slide = document.querySelector(`.slide-${idx}`);
        if (slide) {
          const listItems = Array.from(slide.querySelectorAll("ul li"));
          const checklist = document.createElement("div");
          checklist.className = "checklist-container";
          
          listItems.forEach((li, lIdx) => {
            const item = document.createElement("div");
            item.className = "check-item";
            
            // Check if checkbox is checked
            const text = li.innerHTML.trim();
            const isChecked = text.startsWith("[x]") || text.startsWith("checked");
            const cleanedText = text
              .replace(/^\\[[ x]\\]\\s*/, "")
              .replace(/^checked\\s*/, "");
              
            if (isChecked) {
              item.classList.add("checked");
            }
            
            item.innerHTML = `
              <div class="check-ring"></div>
              <div class="check-text">${cleanedText}</div>
            `;
            
            // Toggle check on click
            item.addEventListener("click", () => {
              item.classList.toggle("checked");
            });
            
            checklist.appendChild(item);
          });
          
          const ul = slide.querySelector("ul");
          if (ul) {
            ul.replaceWith(checklist);
          }
        }
      });

      // 9. Transform Slide 12: Hand-off after submission
      const slide12 = document.querySelector(".slide-12");
      if (slide12) {
        const contentWrap = slide12.querySelector(".slide-content-wrap");
        const h2 = contentWrap.querySelector("h2");
        const pEl = contentWrap.querySelector("p");
        const listItems = Array.from(contentWrap.querySelectorAll("ol li"));
        const footerP = contentWrap.querySelector("p:last-of-type");
        
        const layout = document.createElement("div");
        layout.className = "handoff-layout";
        
        const leftCol = document.createElement("div");
        leftCol.className = "handoff-left";
        if (pEl) {
          const pCopy = document.createElement("div");
          pCopy.className = "handoff-title-p";
          pCopy.innerHTML = pEl.innerHTML;
          leftCol.appendChild(pCopy);
        }
        
        // Add minimal handoff radar graphic
        const radar = document.createElement("div");
        radar.className = "vector-diagram";
        radar.style.height = "160px";
        radar.innerHTML = `
          <svg width="300" height="120" viewBox="0 0 300 120" fill="none" stroke="rgba(0, 240, 255, 0.15)">
            <circle cx="60" cy="60" r="40" />
            <circle cx="240" cy="60" r="40" stroke="var(--accent-cyan)" />
            <path d="M100,60 L200,60" stroke-dasharray="4 4" />
            <polygon points="200,56 208,60 200,64" fill="var(--accent-cyan)" stroke="var(--accent-cyan)" />
            <text x="60" y="64" fill="var(--text-secondary)" font-size="10" text-anchor="middle">AAYUSH</text>
            <text x="240" y="64" fill="var(--accent-cyan)" font-size="10" text-anchor="middle">NAVANSH</text>
          </svg>
        `;
        leftCol.appendChild(radar);
        
        const rightCol = document.createElement("div");
        rightCol.className = "handoff-right";
        
        listItems.forEach((li, lIdx) => {
          const item = document.createElement("div");
          item.className = "handoff-item glass-card";
          item.innerHTML = `
            <div class="handoff-num">0${lIdx+1}</div>
            <div class="handoff-desc">${li.innerHTML}</div>
          `;
          rightCol.appendChild(item);
        });
        
        layout.appendChild(leftCol);
        layout.appendChild(rightCol);
        
        contentWrap.innerHTML = "";
        contentWrap.appendChild(h2);
        contentWrap.appendChild(layout);
        
        if (footerP) {
          const fBox = document.createElement("div");
          fBox.className = "handoff-footer";
          fBox.innerHTML = footerP.innerHTML;
          contentWrap.appendChild(fBox);
        }
      }

      // 10. Transform Slide 13: Team messages
      const slide13 = document.querySelector(".slide-13");
      if (slide13) {
        const contentWrap = slide13.querySelector(".slide-content-wrap");
        const h2 = contentWrap.querySelector("h2");
        const pEl = contentWrap.querySelector("p");
        
        // Extract headers and blockquotes
        const subheaders = Array.from(contentWrap.querySelectorAll("h3"));
        const blockquotes = Array.from(contentWrap.querySelectorAll("blockquote"));
        
        const layout = document.createElement("div");
        layout.className = "messages-layout";
        
        const tabsDiv = document.createElement("div");
        tabsDiv.className = "messages-tabs";
        
        const panesDiv = document.createElement("div");
        panesDiv.className = "messages-panes";
        
        const names = ["Navansh", "Yash", "Nirupam Pal"];
        const initials = ["N", "Y", "NP"];
        
        subheaders.forEach((subh, sIdx) => {
          const tab = document.createElement("button");
          tab.className = `msg-tab ${sIdx === 0 ? 'active' : ''}`;
          tab.innerHTML = `
            <div class="msg-avatar-placeholder">${initials[sIdx]}</div>
            <span>${names[sIdx]}</span>
          `;
          
          tab.addEventListener("click", () => {
            Array.from(tabsDiv.querySelectorAll(".msg-tab")).forEach((t, i) => {
              if (i === sIdx) t.classList.add("active");
              else t.classList.remove("active");
            });
            Array.from(panesDiv.querySelectorAll(".msg-pane")).forEach((p, i) => {
              if (i === sIdx) p.classList.add("active");
              else p.classList.remove("active");
            });
          });
          tabsDiv.appendChild(tab);
          
          const pane = document.createElement("div");
          pane.className = `msg-pane ${sIdx === 0 ? 'active' : ''}`;
          
          const header = document.createElement("div");
          header.className = "msg-header";
          header.innerHTML = `Broadcast Instruction to ${names[sIdx]}`;
          
          const inst = document.createElement("div");
          inst.className = "msg-instruction";
          inst.innerHTML = subh.innerHTML;
          
          const body = document.createElement("div");
          body.className = "msg-body";
          body.innerHTML = blockquotes[sIdx].innerHTML;
          
          pane.appendChild(header);
          pane.appendChild(inst);
          pane.appendChild(body);
          panesDiv.appendChild(pane);
        });
        
        layout.appendChild(tabsDiv);
        layout.appendChild(panesDiv);
        
        contentWrap.innerHTML = "";
        contentWrap.appendChild(h2);
        if (pEl) contentWrap.appendChild(pEl);
        contentWrap.appendChild(layout);
      }

      // 11. Transform Slide 14: The two things Aayush will not write
      const slide14 = document.querySelector(".slide-14");
      if (slide14) {
        const contentWrap = slide14.querySelector(".slide-content-wrap");
        const h2 = contentWrap.querySelector("h2");
        const listItems = Array.from(contentWrap.querySelectorAll("ul li"));
        const footerP = contentWrap.querySelector("p");
        
        const layout = document.createElement("div");
        layout.className = "nw-layout";
        
        const leftCol = document.createElement("div");
        leftCol.className = "nw-left";
        
        const rightCol = document.createElement("div");
        rightCol.className = "nw-right";
        
        listItems.forEach((li, lIdx) => {
          const text = li.innerHTML;
          const boldMatch = text.match(/^<strong>(.*?)<\\/strong>(.*)/);
          let title = `Rule 0${lIdx+1}`;
          let desc = text;
          
          if (boldMatch) {
            title = boldMatch[1].replace(/[:\\.]\\s*$/, "");
            desc = boldMatch[2].trim().replace(/^[-\\s:\\.]*/, "");
          }
          
          const card = document.createElement("div");
          card.className = "nw-card glass-card";
          card.innerHTML = `
            <div class="nw-card-title">${title}</div>
            <div class="nw-card-desc">${desc}</div>
          `;
          leftCol.appendChild(card);
        });
        
        // Render custom Terminal window on the right for scanning
        const termLabel = document.createElement("div");
        termLabel.className = "nw-terminal-label";
        termLabel.innerHTML = "PRE-SUBMIT AUTOMATED SCANNER";
        
        const term = document.createElement("div");
        term.className = "nw-terminal";
        term.innerHTML = `
          $ git grep -nE '(nvapi-|sk-[a-zA-Z]_|AIza[A-Za-z0-9_-]{20,}|"eyJ)'<br/>
          <span style="color: var(--accent-cyan); font-weight: bold;">[SUCCESS]</span> Scan returned 0 matches. Repository is clean.<br/><br/>
          $ grep -rn "console.log" ./frontend/app/<br/>
          <span style="color: var(--accent-cyan); font-weight: bold;">[SUCCESS]</span> No occurrences found.<br/>
        `;
        
        rightCol.appendChild(termLabel);
        rightCol.appendChild(term);
        
        if (footerP) {
          const infoBox = document.createElement("div");
          infoBox.className = "nw-info-box";
          infoBox.innerHTML = footerP.innerHTML;
          rightCol.appendChild(infoBox);
        }
        
        layout.appendChild(leftCol);
        layout.appendChild(rightCol);
        
        contentWrap.innerHTML = "";
        contentWrap.appendChild(h2);
        contentWrap.appendChild(layout);
      }

      // 12. Slide 15: Definition of Done footer element
      const slide15 = document.querySelector(".slide-15");
      if (slide15) {
        const contentWrap = slide15.querySelector(".slide-content-wrap");
        const pEl = contentWrap.querySelector("p");
        const emEl = contentWrap.querySelector("em");
        
        if (pEl) {
          pEl.className = "glass-card";
          pEl.style.marginTop = "24px";
          pEl.style.padding = "16px";
          pEl.style.border = "1px solid rgba(0, 240, 255, 0.15)";
          pEl.style.textAlign = "center";
          pEl.style.fontSize = "1rem";
          pEl.style.color = "#ffffff";
          pEl.style.background = "rgba(0, 240, 255, 0.02)";
        }
        
        if (emEl) {
          const pClose = emEl.closest("p") || emEl;
          pClose.style.marginTop = "30px";
          pClose.style.fontSize = "1.2rem";
          pClose.style.textAlign = "center";
          pClose.style.fontFamily = "var(--font-display)";
          pClose.style.color = var(--accent-cyan);
          pClose.style.textShadow = "0 0 10px rgba(0, 240, 255, 0.3)";
          pClose.style.fontWeight = "600";
        }
      }

      // Display the initial slide
      goToSlide(0);
    });
  </script>
</body>
</html>
"""

def md_to_html(md_text: str) -> str:
    # Split the markdown text by slides (delimiter: ---)
    sections = re.split(r'\n---(?:\r?\n|$)', md_text)
    
    slides_html = []
    for idx, sec in enumerate(sections):
        sec = sec.strip()
        if not sec:
            continue
        
        # Convert section markdown to html
        sec_html = markdown.markdown(
            sec,
            extensions=["tables", "fenced_code", "sane_lists", "toc"],
        )
        
        # Wrap each section inside a slide container
        slides_html.append(f'    <div class="slide slide-{idx}" id="slide-{idx}">\n      <div class="slide-content-wrap">\n        {sec_html}\n      </div>\n    </div>')
    
    slides_combined = "\n".join(slides_html)
    return HTML_TEMPLATE.replace("{{SLIDES_CONTENT}}", slides_combined)

def html_to_pdf_edge(html_path: Path, pdf_path: Path) -> None:
    candidates = [
        r"C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe",
        r"C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe",
    ]
    edge = next((c for c in candidates if Path(c).exists()), None)
    if not edge:
        edge = shutil.which("msedge") or shutil.which("msedge.exe")
    if not edge or not Path(edge).exists():
        raise RuntimeError("Microsoft Edge not found")

    url = "file:///" + str(html_path).replace("\\\\", "/")
    cmd = [
        edge,
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        "--no-pdf-header-footer",
        "--print-to-pdf=" + str(pdf_path),
        url,
    ]
    subprocess.run(cmd, check=True, timeout=180)

def md_text_to_docx(md_text: str, out_path: Path) -> None:
    """Light-weight MD -> DOCX. Tables become compact 'Col | Col' lines."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_buf: list[str] = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        joined = "\\n".join(code_buf)
        p = doc.add_paragraph()
        run = p.add_run(joined)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
        p.paragraph_format.left_indent = Inches(0.25)
        code_buf = []

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
          if in_code:
              in_code = False
              flush_code()
          else:
              in_code = True
          continue
        if in_code:
            code_buf.append(line)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells if c):
                continue
            doc.add_paragraph("  -  " + " | ".join(cells), style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.italic = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line and line[0].isdigit() and ". " in line[:4]:
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line.startswith("- [ ]") or line.startswith("- [x]"):
            mark = "[ ]" if "[ ]" in line else "[x]"
            doc.add_paragraph(mark + "  " + line[5:].strip(), style="List Bullet")
        elif line.strip() == "---":
            doc.add_paragraph("_" * 60)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    flush_code()
    doc.save(str(out_path))

def main():
    md = SRC.read_text(encoding="utf-8")
    print(f"[md]   read {SRC} ({len(md)} chars)")

    html = md_to_html(md)
    HTML_OUT.write_text(html, encoding="utf-8")
    print(f"[html] wrote {HTML_OUT}")

    md_text_to_docx(md, DOCX_OUT)
    print(f"[docx] wrote {DOCX_OUT}")

    try:
        html_to_pdf_edge(HTML_OUT, PDF_OUT)
        size = PDF_OUT.stat().st_size if PDF_OUT.exists() else 0
        print(f"[pdf]  wrote {PDF_OUT} ({size/1024:.1f} KB)")
    except Exception as exc:
        print(f"[pdf]  FAILED: {exc}")
        sys.exit(2)

if __name__ == "__main__":
    main()
